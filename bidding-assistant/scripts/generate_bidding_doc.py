#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
生成标准格式投标文件 .docx

参考 references/standard_structure.md 的九章结构。
绝不包含中标概率、匹配度评分等自评内容（参考 references/forbidden_content.md）。

依赖: python-docx
安装: pip install python-docx

用法:
    python generate_bidding_doc.py [--data DATA_DIR]
"""

import os
import sys
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml


# ==================== 工具函数 ====================

def set_east_asian_font(run, font_name="宋体"):
    """安全设置东亚字体"""
    rPr = run.element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = parse_xml(f'<w:rFonts {nsdecls("w")} w:eastAsia="{font_name}"/>')
        rPr.append(rFonts)
    else:
        rFonts.set(qn('w:eastAsia'), font_name)


def add_table_shading(cell, color):
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color}"/>')
    cell._tc.get_or_add_tcPr().append(shading)


def set_cell(cell, text, bold=False, align=WD_ALIGN_PARAGRAPH.LEFT,
             size=10, color=None, font_name="宋体"):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = align
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.name = font_name
    run.bold = bold
    set_east_asian_font(run, font_name)
    if color:
        run.font.color.rgb = color


def make_header(table, headers, col_widths=None):
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        set_cell(cell, h, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER,
                 size=10, color=RGBColor(0xFF, 0xFF, 0xFF), font_name="黑体")
        add_table_shading(cell, "1a3a5c")
    if col_widths:
        for row in table.rows:
            for i, w in enumerate(col_widths):
                row.cells[i].width = Cm(w)


def add_body(doc, text, indent=True):
    p = doc.add_paragraph()
    if indent:
        p.paragraph_format.first_line_indent = Cm(0.74)
    run = p.add_run(text)
    run.font.size = Pt(12)
    run.font.name = "宋体"
    set_east_asian_font(run, "宋体")
    return p


def add_bullet(doc, text):
    p = doc.add_paragraph(f"■ {text}")
    p.paragraph_format.left_indent = Cm(1)
    for run in p.runs:
        run.font.size = Pt(12)
    return p


def add_chapter_title(doc, title):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(title)
    run.font.size = Pt(22)
    run.font.name = "黑体"
    set_east_asian_font(run, "黑体")
    run.bold = True
    doc.add_paragraph()


def make_data_table(doc, headers, data, col_widths=None):
    t = doc.add_table(rows=len(data) + 1, cols=len(headers))
    t.style = 'Table Grid'
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    make_header(t, headers, col_widths)
    for i, row_data in enumerate(data):
        row = t.rows[i + 1]
        for j, val in enumerate(row_data):
            set_cell(row.cells[j], val)
        if i % 2 == 0:
            for cell in row.cells:
                add_table_shading(cell, "f8f9fb")
    return t


# ==================== 生成投标文件 ====================

def generate_bidding_doc(output_path, project_data=None):
    """生成标准格式投标文件

    project_data: dict, 可选，用于填充项目特定信息
        {
            "project_name": "项目名称",
            "bidder_name": "投标人全称",
            "purchaser_name": "采购人全称",
            "bid_number": "招标编号",
            "delivery_days": 45,
            "free_maintenance_months": 12,
            "training_admin_sessions": 3,
            "training_user_sessions": 4,
            "response_hours": 4,
            ...
        }
    """
    if project_data is None:
        project_data = {}

    PROJ = project_data.get("project_name", "【项目名称】")
    BIDDER = project_data.get("bidder_name", "【投标人全称】")
    PURCHASER = project_data.get("purchaser_name", "【采购人全称】")
    BID_NO = project_data.get("bid_number", "【待填入招标编号】")
    DAYS = project_data.get("delivery_days", 45)
    MONTHS = project_data.get("free_maintenance_months", 12)
    ADMIN_SESSIONS = project_data.get("training_admin_sessions", 3)
    USER_SESSIONS = project_data.get("training_user_sessions", 4)
    RESP_HOURS = project_data.get("response_hours", 4)

    doc = Document()

    # 全局样式
    style = doc.styles['Normal']
    style.font.name = '宋体'
    style.font.size = Pt(12)
    set_east_asian_font(style.font, '宋体')
    style.paragraph_format.line_spacing = 1.5

    # ============ 封面 ============
    for _ in range(6):
        doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("投 标 文 件")
    r.font.size = Pt(18)
    r.font.name = "黑体"
    set_east_asian_font(r, "黑体")
    r.font.color.rgb = RGBColor(0x1a, 0x3a, 0x5c)

    doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(PROJ)
    r.font.size = Pt(28)
    r.font.name = "黑体"
    set_east_asian_font(r, "黑体")
    r.font.color.rgb = RGBColor(0x1a, 0x3a, 0x5c)
    r.bold = True

    doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("（技术商务部分）")
    r.font.size = Pt(16)

    for _ in range(4):
        doc.add_paragraph()

    for label, value in [
        ("采 购 人：", PURCHASER),
        ("招标编号：", BID_NO),
        ("投标人：", BIDDER),
        ("地    址：", "【地址】"),
        ("联系电话：", "【电话】"),
        ("日    期：", "2026年   月   日"),
    ]:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r1 = p.add_run(label)
        r1.font.size = Pt(14)
        r2 = p.add_run(value)
        r2.font.size = Pt(14)
        r2.underline = True

    doc.add_page_break()

    # ============ 第一章 投标函 ============
    add_chapter_title(doc, "投 标 函")
    doc.add_paragraph()

    add_body(doc, f"致：{PURCHASER}", indent=False)
    doc.add_paragraph()

    add_body(doc, f'根据贵方"{PROJ}"（招标编号：{BID_NO}）的招标文件，我方经认真研究招标文件的全部内容，愿意以本投标函附录所列的投标总价，按照招标文件的要求，承担本项目的建设实施工作。')

    doc.add_paragraph()
    add_body(doc, "一、我方承诺：")

    commitments = [
        f"（1）本投标文件的有效期为自投标截止日起90个日历天。在此期间，本投标文件对我方具有法律约束力。",
        f"（2）我方承诺在合同签订后{DAYS}个自然日内完成系统的部署、配置、试运行和首轮培训。",
        f"（3）我方承诺提供不少于{MONTHS}个月的免费运维服务（含故障修复、使用指导、版本升级）。",
        f"（4）我方承诺在收到故障通知后{RESP_HOURS}小时内响应并提供解决方案。",
        f"（5）我方承诺提供至少{ADMIN_SESSIONS}场管理员深度培训和{USER_SESSIONS}场业务用户操作培训。",
        f"（6）我方承诺严格履行保密义务，不将项目资料用于其他用途。",
        f"（7）我方承诺所有投标内容真实有效，如存在虚假，愿承担一切法律责任。",
    ]
    for c in commitments:
        p = doc.add_paragraph()
        p.paragraph_format.first_line_indent = Cm(1.5)
        r = p.add_run(c)
        r.font.size = Pt(12)

    doc.add_paragraph()
    add_body(doc, "二、我方理解，贵方有权接受或拒绝任何投标，且无义务对任何投标做出解释。")

    doc.add_paragraph()
    doc.add_paragraph()

    for label, value in [
        ("投标人（盖章）：", BIDDER),
        ("法定代表人或授权代表（签字）：", ""),
        ("日期：", "2026年   月   日"),
    ]:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        r1 = p.add_run(label)
        r1.font.size = Pt(12)
        if value:
            r2 = p.add_run(value)
            r2.font.size = Pt(12)
            r2.underline = True

    doc.add_page_break()

    # ============ 投标函附录 ============
    add_chapter_title(doc, "投标函附录")
    doc.add_paragraph()

    make_data_table(
        doc,
        ["序号", "项目名称", "招标文件要求", "投标响应"],
        [
            ("1", "投标有效期", "自投标截止日起90个日历天", "完全响应，有效期90个日历天"),
            ("2", "交付周期", f"合同签订后{DAYS}个自然日", f"完全响应，承诺{DAYS}个自然日内完成"),
            ("3", f"质保期/免费运维", f"不少于{MONTHS}个月", f"完全响应，提供{MONTHS}个月免费运维"),
            ("4", "故障响应", f"重大故障{RESP_HOURS}小时内响应", f"完全响应，{RESP_HOURS}小时内响应"),
            ("5", "培训要求", f"至少2场管理员培训+3场业务用户培训",
             f"完全响应，提供{ADMIN_SESSIONS}场管理员培训+{USER_SESSIONS}场业务用户培训"),
            ("6", "付款方式", "按招标文件要求", "完全响应"),
            ("7", "保密义务", "签署保密协议", "完全响应，签署严格保密协议"),
            ("8", "其他实质性要求", "按招标文件要求", "完全响应，无偏离"),
        ],
        col_widths=[1.5, 3, 5, 5]
    )

    doc.add_page_break()

    # ============ 第二章 法定代表人证明及授权书 ============
    add_chapter_title(doc, "法定代表人身份证明书")
    doc.add_paragraph()
    doc.add_paragraph()
    add_body(doc, "兹证明：")
    add_body(doc, "【法定代表人姓名】，性别：【男/女】，身份证号码：【身份证号码】，系" + BIDDER + "的法定代表人，担任公司【职务】职务。")
    doc.add_paragraph()
    add_body(doc, "特此证明。")

    for _ in range(3):
        doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = p.add_run(f"投标人（盖章）：{BIDDER}")
    r.font.size = Pt(12)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = p.add_run("日    期：2026年   月   日")
    r.font.size = Pt(12)

    doc.add_paragraph()
    doc.add_paragraph()

    add_chapter_title(doc, "授权委托书")
    doc.add_paragraph()
    doc.add_paragraph()
    add_body(doc, f"本人【法定代表人姓名】系{BIDDER}的法定代表人，现委托【被授权人姓名】作为我方代理人，就\"{PROJ}\"（招标编号：{BID_NO}）的投标及合同签订事宜，以我方名义全权处理。")
    doc.add_paragraph()
    add_body(doc, "代理人的代理权限包括但不限于：签署投标文件、参加开标会议、进行商务谈判、签署合同及相关文件。代理人在授权范围内所签署的一切文件，我方均予认可。")
    doc.add_paragraph()
    add_body(doc, "委托期限：自本委托书签署之日起至本项目合同履行完毕之日止。")

    doc.add_paragraph()
    doc.add_paragraph()

    for label in [
        "法定代表人（签字）：",
        "身份证号码：【身份证号码】",
        "日    期：2026年   月   日",
    ]:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        r = p.add_run(label)
        r.font.size = Pt(12)

    doc.add_page_break()

    # ============ 第三章 资格证明文件 ============
    add_chapter_title(doc, "资格证明文件")

    doc.add_heading("一、资格证明文件清单", level=1)

    qual_headers = ["序号", "文件名称", "要求", "提供情况", "页码"]
    qual_data = [
        ("1", "营业执照（副本）", "复印件加盖公章", "已提供", "见附件"),
        ("2", "ISO 9001质量管理体系认证证书", "有效期内复印件", "已提供", "见附件"),
        ("3", "ISO 27001信息安全管理体系认证证书", "有效期内复印件", "已提供", "见附件"),
        ("4", "CMMI认证证书", "有效期内复印件", "申报中（详见说明）", "见附件"),
        ("5", "软件著作权登记证书", "复印件加盖公章", "已提供", "见附件"),
        ("6", "近三年财务审计报告", "会计师事务所出具", "已提供", "见附件"),
        ("7", "近3个月依法缴纳税收证明", "税务部门出具", "已提供", "见附件"),
        ("8", "近3个月社会保障资金缴纳证明", "社保部门出具", "已提供", "见附件"),
        ("9", "无重大违法记录声明", "加盖公章", "已提供", "见下"),
    ]
    make_data_table(doc, qual_headers, qual_data, col_widths=[1.2, 4, 3, 3, 1.5])

    doc.add_paragraph()

    doc.add_heading("二、关于CMMI认证的说明", level=2)
    add_body(doc, f"{BIDDER}高度重视软件研发过程的规范化管理，已按照CMMI-3级标准建立了完整的研发过程管理体系，涵盖需求管理、项目策划、度量与分析、过程与产品质量保证、配置管理等关键过程域。公司目前处于CMMI-3级认证申报阶段，预计将于近期获得正式评估。在此期间，公司已通过ISO 9001质量管理体系和ISO 27001信息安全管理体系认证，能够确保项目交付质量。")

    doc.add_paragraph()

    doc.add_heading("三、无重大违法记录声明", level=2)
    add_body(doc, "本公司郑重声明：近三年内，本公司及法定代表人无重大违法记录，未受到过刑事处罚或与经营相关的行政处罚。本公司所提交的所有投标文件内容真实、完整、有效，不存在任何虚假陈述。如有不实，愿承担一切法律责任。")

    doc.add_page_break()

    # ============ CHAPTER PLACEHOLDER ============
    # 后续章节（第四-九章 + 签署页）请根据项目实际情况填充
    # 参考 references/standard_structure.md 了解完整结构

    add_chapter_title(doc, "技术方案")
    add_body(doc, "【请根据具体项目填充技术方案内容，至少包含：项目概述与需求理解、总体架构设计、功能模块设计、AI能力方案、数据安全与合规】")

    doc.add_page_break()

    add_chapter_title(doc, "项目实施方案")
    add_body(doc, "【请根据具体项目填充实施方案内容，至少包含：实施里程碑计划、项目团队配置、质量管理方案】")

    doc.add_page_break()

    add_chapter_title(doc, "培训与售后服务方案")
    add_body(doc, "【请根据具体项目填充培训与售后服务内容】")

    doc.add_page_break()

    add_chapter_title(doc, "商务响应表")
    add_body(doc, "【请根据具体项目填充商务响应内容，含报价清单】")

    doc.add_page_break()

    add_chapter_title(doc, "业绩与案例")
    add_body(doc, "【请根据具体项目填充业绩与案例内容】")

    doc.add_page_break()

    add_chapter_title(doc, "其他材料")
    add_body(doc, "【请根据具体项目填充保密承诺、知识产权声明、风险控制等】")

    doc.add_page_break()

    # ============ 签署页 ============
    for _ in range(6):
        doc.add_paragraph()

    add_chapter_title(doc, "投标文件签署页")
    doc.add_paragraph()
    doc.add_paragraph()
    add_body(doc, "本单位郑重承诺：本投标文件所有内容真实、有效，如有虚假，愿承担一切法律责任。")

    for _ in range(4):
        doc.add_paragraph()

    for label, value in [
        ("投标人（盖章）：", BIDDER),
        ("法定代表人或授权代表（签字）：", ""),
        ("日    期：", "2026年   月   日"),
    ]:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        r1 = p.add_run(label)
        r1.font.size = Pt(14)
        if value:
            r2 = p.add_run(value)
            r2.font.size = Pt(14)
            r2.underline = True

    # 保存
    doc.save(output_path)
    print(f"投标文件已生成: {output_path}")
    return output_path


if __name__ == "__main__":
    out = sys.argv[1] if len(sys.argv) > 1 else "投标文件.docx"
    generate_bidding_doc(out)
